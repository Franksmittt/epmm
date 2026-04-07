"""
Build Miwesu Master Knowledge: one large txt from all Miwesu-relevant
content (farm, hunting, lodge, target audience, firewood, strategies)
in Game Reserve and Hunting, Firewood, and Miwesu_Farm.
"""
from pathlib import Path
import sys

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

DOWNLOADS = Path(r"C:\Users\User1\Downloads")
GAME_RESERVE = DOWNLOADS / "Game Reserve and Hunting"
FIREWOOD = DOWNLOADS / "Firewood"
MIWESU_FARM = DOWNLOADS / "Miwesu_Farm"
MIWESU_WOOD = DOWNLOADS / "Miwesu_Wood_Creatives"
OUTPUT_FILE = MIWESU_FARM / "Miwesu_Master_Knowledge.txt"

# Sections to include: (folder, description)
SOURCES = [
    (GAME_RESERVE, "Game Reserve and Hunting (hunting, lodge, game farm, target audience, SEO)"),
    (FIREWOOD, "Firewood / Braai Wood (Miwesu firewood business strategies)"),
    (MIWESU_FARM, "Miwesu Farm (existing docs)"),
    (MIWESU_WOOD, "Miwesu Wood Creatives (firewood creatives / ads)"),
]

def read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception as e:
        return f"[Error reading file: {e}]"

def read_docx(path: Path) -> str:
    if DocxDocument is None:
        return "[python-docx not installed; skipping .docx]"
    try:
        doc = DocxDocument(path)
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as e:
        return f"[Error reading docx: {e}]"

def main():
    out_lines = []
    out_lines.append("=" * 80)
    out_lines.append("MIWESU MASTER KNOWLEDGE")
    out_lines.append("Farm • Hunting • Lodge • Target Audience • Firewood • Strategies & Research")
    out_lines.append("=" * 80)
    out_lines.append("")

    for folder, description in SOURCES:
        if not folder.exists():
            continue
        # Text files
        for path in sorted(folder.rglob("*.txt")):
            if path.name.startswith(".") or "Miwesu_Master" in path.name or "build_miwesu" in path.name:
                continue
            content = read_txt(path)
            if not content or len(content) < 20:
                continue
            rel = path.relative_to(folder)
            out_lines.append("")
            out_lines.append("-" * 80)
            out_lines.append(f"SOURCE: {rel} ({description})")
            out_lines.append("-" * 80)
            out_lines.append("")
            out_lines.append(content)
            out_lines.append("")

        # Word documents
        for path in sorted(folder.rglob("*.docx")):
            if path.name.startswith("~"):
                continue
            content = read_docx(path)
            if not content or len(content) < 20:
                continue
            rel = path.relative_to(folder)
            out_lines.append("")
            out_lines.append("-" * 80)
            out_lines.append(f"SOURCE: {rel} ({description})")
            out_lines.append("-" * 80)
            out_lines.append("")
            out_lines.append(content)
            out_lines.append("")

    full = "\n".join(out_lines)
    OUTPUT_FILE.write_text(full, encoding="utf-8")
    print(f"Written: {OUTPUT_FILE}")
    print(f"Size: {len(full):,} chars, {len(full)/1024:.1f} KB")
    return 0

if __name__ == "__main__":
    sys.exit(main())
